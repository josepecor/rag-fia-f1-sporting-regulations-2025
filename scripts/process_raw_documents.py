# ====================================
# process_f1_documents.py
# Procesa documentos F1: convierte DOCX a TXT y reemplaza referencias [archivo]
# ====================================

import os
import yaml
import re
from pathlib import Path
from typing import Dict, List, Optional, Set
import sys
import shutil
import argparse

# Para leer DOCX
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx no instalado. Para convertir DOCX ejecuta: pip install python-docx")

class F1DocumentProcessor:
    """Procesa documentos F1 según índice YAML"""
    
    def __init__(
        self, 
        yaml_path: str = None,
        raw_dir: str = "../data/raw",
        output_dir: str = "../data/processed"
    ):
        # Obtener directorio del script para construir rutas absolutas
        script_dir = Path(__file__).parent.absolute()
        project_root = script_dir.parent  # Subir un nivel desde scripts/
        
        self.yaml_path = yaml_path
        
        # Convertir rutas relativas a absolutas basadas en project_root
        if not os.path.isabs(raw_dir):
            self.raw_dir = str(project_root / raw_dir.lstrip('../'))
        else:
            self.raw_dir = raw_dir
            
        if not os.path.isabs(output_dir):
            self.output_dir = str(project_root / output_dir.lstrip('../'))
        else:
            self.output_dir = output_dir
            
        self.index_data = None
        self.file_cache = {}
        
        print(f"🔧 Directorio del proyecto: {project_root}")
        print(f"🔧 Buscando archivos raw en: {self.raw_dir}")
        print(f"🔧 Salida en: {self.output_dir}")
    
    def find_yaml_file(self) -> Optional[str]:
        """Busca automáticamente el archivo YAML"""
        
        print("\n🔍 Buscando archivo YAML...")
        
        # Obtener project_root
        script_dir = Path(__file__).parent.absolute()
        project_root = script_dir.parent
        
        # Buscar en data/
        search_paths = [
            project_root / 'data',
            script_dir / 'data',
            Path.cwd() / 'data'
        ]
        
        patterns = [
            'sporting_regulations_2025_index.yaml'
        ]
        
        found_files = []
        
        for search_dir in search_paths:
            if not search_dir.exists():
                continue
            
            for pattern in patterns:
                if '*' in pattern:
                    import glob
                    matches = glob.glob(str(search_dir / pattern))
                    found_files.extend(matches)
                else:
                    filepath = search_dir / pattern
                    if filepath.exists():
                        found_files.append(str(filepath))
        
        found_files = list(set(found_files))
        
        if not found_files:
            print("❌ No se encontraron archivos YAML")
            return None
        
        print(f"\n📋 Archivos YAML encontrados:")
        for i, f in enumerate(found_files, 1):
            size = os.path.getsize(f)
            print(f"   {i}. {f} ({size} bytes)")
        
        if len(found_files) == 1:
            selected = found_files[0]
            print(f"\n✅ Seleccionado: {selected}")
            return selected
        
        print(f"\n¿Cuál archivo deseas usar? (1-{len(found_files)}, 0 para cancelar): ")
        try:
            choice = int(input().strip())
            if choice == 0:
                return None
            if 1 <= choice <= len(found_files):
                selected = found_files[choice - 1]
                print(f"✅ Seleccionado: {selected}")
                return selected
            else:
                print("❌ Opción inválida")
                return None
        except (ValueError, KeyboardInterrupt):
            print("\n❌ Cancelado")
            return None
    
    def load_yaml_index(self):
        """Carga el índice YAML"""
        
        if self.yaml_path is None:
            self.yaml_path = self.find_yaml_file()
            if self.yaml_path is None:
                raise FileNotFoundError("No se pudo encontrar archivo YAML")
        
        if not os.path.exists(self.yaml_path):
            print(f"⚠️  Ruta no encontrada: {self.yaml_path}")
            self.yaml_path = self.find_yaml_file()
            if self.yaml_path is None:
                raise FileNotFoundError("No se pudo encontrar archivo YAML")
        
        print(f"\n📂 Cargando: {os.path.abspath(self.yaml_path)}")
        
        try:
            with open(self.yaml_path, 'r', encoding='utf-8') as f:
                content = f.read()
            print(f"✅ Archivo leído ({len(content)} caracteres)")
        except UnicodeDecodeError:
            with open(self.yaml_path, 'r', encoding='latin-1') as f:
                content = f.read()
        
        try:
            self.index_data = yaml.safe_load(content)
            
            if self.index_data is None:
                raise ValueError("Archivo YAML vacío")
            
            print(f"✅ YAML parseado")
            self._validate_structure()
            
        except yaml.YAMLError as e:
            print(f"\n❌ ERROR parseando YAML: {e}")
            if hasattr(e, 'problem_mark'):
                mark = e.problem_mark
                print(f"📍 Línea {mark.line + 1}, columna {mark.column + 1}")
            raise
    
    def _validate_structure(self):
        """Valida estructura del YAML"""
        
        print(f"🔍 Validando estructura...")
        
        if 'document' not in self.index_data:
            raise ValueError("Falta 'document'")
        
        if 'structure' not in self.index_data:
            raise ValueError("Falta 'structure'")
        
        structure = self.index_data['structure']
        
        if 'articles' not in structure:
            raise ValueError("Falta 'articles'")
        
        if 'appendices' not in structure:
            structure['appendices'] = []
        
        num_articles = len(structure['articles'])
        num_appendices = len(structure['appendices'])
        
        print(f"✅ {num_articles} artículos, {num_appendices} apéndices")
    
    def read_docx(self, filepath: str) -> Optional[str]:
        """Convierte DOCX a texto plano"""
        
        if not DOCX_AVAILABLE:
            print(f"   ⚠️  python-docx no disponible, no se puede leer DOCX")
            return None
        
        try:
            doc = DocxDocument(filepath)
            
            # Extraer todo el texto
            full_text = []
            
            # Párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            return '\n\n'.join(full_text)
            
        except Exception as e:
            print(f"   ❌ Error leyendo DOCX: {e}")
            return None
    
    def read_file(self, filename: str) -> Optional[str]:
        """Lee un archivo del directorio raw"""
        
        if filename in self.file_cache:
            return self.file_cache[filename]
        
        filepath = os.path.join(self.raw_dir, filename)
        
        print(f"   🔍 Intentando leer: {filepath}")
        print(f"   📂 Ruta absoluta: {os.path.abspath(filepath)}")
        print(f"   ❓ Existe: {os.path.exists(filepath)}")
        
        if not os.path.exists(filepath):
            print(f"   ⚠️  No encontrado: {filename}")
            # Mostrar contenido del directorio para debug
            if os.path.exists(self.raw_dir):
                print(f"   📁 Contenido de {self.raw_dir}:")
                for f in os.listdir(self.raw_dir)[:10]:  # Primeros 10 archivos
                    print(f"      - {f}")
            return None
        
        try:
            # TXT y MD
            if filename.endswith(('.txt', '.md')):
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # DOCX - convertir a texto
            elif filename.endswith('.docx'):
                print(f"   📄 Convirtiendo DOCX: {filename}")
                content = self.read_docx(filepath)
                if content is None:
                    return None
            
            # CSV y JSON
            elif filename.endswith(('.csv', '.json')):
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            else:
                print(f"   ⚠️  Formato no soportado: {filename}")
                return None
            
            self.file_cache[filename] = content
            return content
            
        except UnicodeDecodeError:
            try:
                with open(filepath, 'r', encoding='latin-1') as f:
                    content = f.read()
                self.file_cache[filename] = content
                return content
            except Exception as e:
                print(f"   ❌ Error: {e}")
                return None
        except Exception as e:
            print(f"   ❌ Error: {e}")
            return None
    
    def find_file_references(self, text: str) -> Set[str]:
        """
        Encuentra referencias a archivos en formato [nombre_archivo.ext]
        Ejemplo: [SR2025_A06_points.csv]
        """
        # Patrón para encontrar [archivo.ext]
        pattern = r'\[([A-Za-z0-9_\-\.]+\.(txt|md|csv|json|docx))\]'
        matches = re.findall(pattern, text)
        return set([match[0] for match in matches])
    
    def replace_file_references(self, text: str, depth: int = 0, max_depth: int = 5) -> str:
        """
        Reemplaza referencias a archivos [archivo.ext] con su contenido
        
        Args:
            text: Texto con referencias
            depth: Nivel de recursión actual
            max_depth: Máximo nivel de recursión
        
        Returns:
            Texto con referencias reemplazadas
        """
        
        if depth >= max_depth:
            print(f"   ⚠️  Máxima profundidad alcanzada ({max_depth})")
            return text
        
        # Encontrar referencias
        file_refs = self.find_file_references(text)
        
        if not file_refs:
            return text
        
        # Reemplazar cada referencia
        for ref_file in file_refs:
            print(f"   {'  ' * depth}🔄 [{ref_file}]")
            
            # Leer contenido del archivo referenciado
            sub_content = self.read_file(ref_file)
            
            if sub_content is None:
                print(f"   {'  ' * depth}⚠️  No se pudo leer, manteniendo referencia")
                continue
            
            # Reemplazar referencia en el texto
            ref_pattern = r'\[' + re.escape(ref_file) + r'\]'
            text = re.sub(ref_pattern, sub_content, text)
        
        # Verificar si hay nuevas referencias (recursión)
        new_refs = self.find_file_references(text)
        
        if new_refs and depth < max_depth - 1:
            print(f"   {'  ' * depth}🔁 Nuevas referencias encontradas")
            text = self.replace_file_references(text, depth + 1, max_depth)
        
        return text
    
    def process_article(self, article: Dict) -> bool:
        """Procesa un artículo - SIN CABECERAS"""
        
        article_num = article.get('article', '?')
        title = article.get('title', 'Unknown')
        filename = article.get('file')
        
        if not filename:
            print(f"\n❌ A{article_num}: sin 'file'")
            return False
        
        print(f"\n📄 A{article_num}: {title}")
        
        # Leer contenido principal
        content = self.read_file(filename)
        
        if content is None:
            return False
        
        # Buscar y reemplazar referencias [archivo.ext]
        file_refs = self.find_file_references(content)
        
        if file_refs:
            print(f"   🔍 {len(file_refs)} referencias a archivos")
            for ref in file_refs:
                print(f"      - [{ref}]")
            content = self.replace_file_references(content)
        else:
            print(f"   ℹ️  Sin referencias")
        
        # Generar nombre de archivo
        safe_title = re.sub(r'[^\w\s-]', '', title.lower())
        safe_title = re.sub(r'[-\s]+', '_', safe_title)
        output_filename = f"SR2025_A{article_num:02d}_{safe_title}.txt"
        output_path = os.path.join(self.output_dir, output_filename)
        
        try:
            # Guardar SIN CABECERAS - solo el contenido
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            print(f"   ✅ {output_filename}")
            return True
            
        except Exception as e:
            print(f"   ❌ {e}")
            return False
    
    def process_appendix(self, appendix: Dict) -> bool:
        """Procesa un apéndice - SIN CABECERAS"""
        
        appendix_num = appendix.get('appendix', '?')
        title = appendix.get('title', 'Unknown')
        filename = appendix.get('file')
        
        if not filename:
            print(f"\n❌ APP{appendix_num}: sin 'file'")
            return False
        
        print(f"\n📋 APP{appendix_num}: {title}")
        
        # Leer contenido (incluyendo DOCX que se convierte a TXT)
        content = self.read_file(filename)
        
        if content is None:
            return False
        
        # Buscar y reemplazar referencias [archivo.ext]
        file_refs = self.find_file_references(content)
        
        if file_refs:
            print(f"   🔍 {len(file_refs)} referencias")
            for ref in file_refs:
                print(f"      - [{ref}]")
            content = self.replace_file_references(content)
        
        # Generar nombre de archivo - SIEMPRE .txt
        safe_title = re.sub(r'[^\w\s-]', '', title.lower())
        safe_title = re.sub(r'[-\s]+', '_', safe_title)[:50]
        output_filename = f"SR2025_APP{appendix_num:02d}_{safe_title}.txt"
        output_path = os.path.join(self.output_dir, output_filename)
        
        try:
            # Guardar SIN CABECERAS - solo el contenido
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            print(f"   ✅ {output_filename}")
            return True
            
        except Exception as e:
            print(f"   ❌ {e}")
            return False
    
    def process_all(self):
        """Procesa todos los documentos"""
        
        print(f"\n{'='*100}")
        print(f"🚀 PROCESAMIENTO DOCUMENTOS F1")
        print(f"{'='*100}\n")
        
        self.load_yaml_index()
        
        os.makedirs(self.output_dir, exist_ok=True)
        print(f"\n📁 Salida: {os.path.abspath(self.output_dir)}")
        
        # Artículos
        print(f"\n{'='*100}")
        print(f"📚 ARTÍCULOS")
        print(f"{'='*100}")
        
        articles = self.index_data['structure']['articles']
        success_articles = 0
        
        for article in articles:
            if self.process_article(article):
                success_articles += 1
        
        # Apéndices
        print(f"\n{'='*100}")
        print(f"📎 APÉNDICES")
        print(f"{'='*100}")
        
        appendices = self.index_data['structure']['appendices']
        success_appendices = 0
        
        for appendix in appendices:
            if self.process_appendix(appendix):
                success_appendices += 1
        
        # Resumen
        total_success = success_articles + success_appendices
        total_items = len(articles) + len(appendices)
        
        print(f"\n{'='*100}")
        print(f"📊 RESUMEN")
        print(f"{'='*100}")
        print(f"\n✅ Procesados: {total_success}/{total_items}")
        print(f"   Artículos: {success_articles}/{len(articles)}")
        print(f"   Apéndices: {success_appendices}/{len(appendices)}")
        print(f"\n📁 Archivos en: {os.path.abspath(self.output_dir)}")
        print(f"{'='*100}\n")


def clean_directory(directory: str):
    """Elimina y recrea un directorio"""
    
    # Convertir a ruta absoluta basada en project_root
    script_dir = Path(__file__).parent.absolute()
    project_root = script_dir.parent
    
    if not os.path.isabs(directory):
        directory = str(project_root / directory)
    
    if os.path.exists(directory):
        shutil.rmtree(directory)
        print(f"✅ {directory} eliminado")
    
    os.makedirs(directory)
    print(f"✅ {directory} recreado vacío")

def main():
    
    clean_directory("data/processed")
    clean_directory("models/best_rag_model")
    clean_directory("outputs/results")
    
    parser = argparse.ArgumentParser(description='Procesa documentos F1 a TXT')
    
    parser.add_argument(
        '--yaml',
        default=None,
        help='Ruta al YAML (búsqueda automática si no se especifica)'
    )
    
    parser.add_argument(
        '--raw-dir',
        default='../data/raw',
        help='Directorio con archivos originales (default: ../data/raw)'
    )
    
    parser.add_argument(
        '--output-dir',
        default='../data/processed',
        help='Directorio de salida (default: ../data/processed)'
    )
    
    args = parser.parse_args()
    
    processor = F1DocumentProcessor(
        yaml_path=args.yaml,
        raw_dir=args.raw_dir,
        output_dir=args.output_dir
    )
    
    try:
        processor.process_all()
        print(f"✅ ¡Completado!")
        return 0
        
    except KeyboardInterrupt:
        print(f"\n\n⚠️  Cancelado")
        return 1
        
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    sys.exit(main())